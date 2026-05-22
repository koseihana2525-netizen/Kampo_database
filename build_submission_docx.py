#!/usr/bin/env python3
"""投稿用docx生成: 原稿v3 Markdown → 図入りdocx"""

import os, re, sys
from docx import Document
from docx.shared import Pt, Mm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE, 'manuscript', 'manuscript_v3.md')
FIG_DIR = os.path.join(BASE, 'submission', 'figures')
OUT_FULL = os.path.join(BASE, 'submission', '01_manuscript.docx')
OUT_TEXT = os.path.join(BASE, 'submission', '02_manuscript_text_only.docx')

# 図ファイルマッピング
FIG_MAP = {
    'Figure 1': ('Figure1_timeseries.png', 'Figure 1. kampo群およびpubmed_kampo群における認知ギャップ比率の経年推移（5年単位）'),
    'Figure 2': ('Figure2_cross_language.png', 'Figure 2. 言語間認知ギャップ：処方別の認知ギャップ比率の日英比較'),
    'Figure 3': ('Figure3_categories.png', 'Figure 3. 処方言及論文における漢方思考概念カテゴリ別出現率の日英比較'),
    'Figure 4': ('Figure4_quadrant_trend.png', 'Figure 4. 4象限分類の年代別推移（kampo群およびpubmed_kampo群）'),
    'Figure 5': ('Figure5_pubtype.png', 'Figure 5. 論文タイプ別認知ギャップ比率（kampo群 vs pubmed群）'),
}

def read_markdown():
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        return f.read()

def build_docx(include_figures=True):
    doc = Document()

    # ページ設定: A4, マージン30mm
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(30)
    section.left_margin = Mm(30)
    section.right_margin = Mm(30)

    # ページ番号（右上）
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run()
    fld_xml = '<w:fldSimple {} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'.format(
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')
    from lxml import etree
    fld_elem = etree.fromstring(fld_xml)
    run._element.append(fld_elem)

    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)
    font.name = 'Yu Mincho'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Mincho')
    pf = style.paragraph_format
    pf.line_spacing = Pt(24)  # ダブルスペース
    pf.space_after = Pt(0)

    text = read_markdown()
    lines = text.split('\n')

    # 原稿の構造をパース
    # タイトルページ
    title_jp = lines[0].lstrip('# ').strip()
    authors = lines[2].strip()
    affil1 = lines[4].strip() if len(lines) > 4 else ''
    affil2 = lines[5].strip() if len(lines) > 5 else ''

    # --- タイトルページ ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(24)
    run = p.add_run(title_jp)
    run.font.size = Pt(16)
    run.bold = True
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    # 英文タイトル
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(24)
    run = p.add_run('Discrepancy Between Formula Citation and Traditional Diagnostic Reasoning\nin Kampo Literature: A Text-Mining Analysis of 10,535 Articles')
    run.font.size = Pt(12)
    run.italic = True
    run.font.name = 'Times New Roman'

    # 著者
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authors)
    run.font.size = Pt(12)

    # 所属
    for aff in [affil1, affil2]:
        if aff and not aff.startswith('---'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(aff)
            run.font.size = Pt(10)

    # 責任著者
    p = doc.add_paragraph()
    p.space_before = Pt(24)
    run = p.add_run('責任著者：華岡晃生')
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run('Email: ★（確定後に記入）')
    run.font.size = Pt(10)

    # Running title
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run('Running title: 漢方文献の認知ギャップ')
    run.font.size = Pt(10)

    doc.add_page_break()

    # --- 本文の処理 ---
    # セクション開始行を特定
    body_start = None
    ref_start = None
    for i, line in enumerate(lines):
        if line.strip() == '## 緒言' and body_start is None:
            body_start = i
        if line.strip() == '## 引用文献':
            ref_start = i

    if body_start is None or ref_start is None:
        print("ERROR: cannot find sections")
        return

    # 要旨・Abstract・キーワードをmarkdownから抽出
    ja_abstract_lines = []
    en_abstract_lines = []
    ja_keywords = ''
    en_keywords = ''
    section = None
    for i, line in enumerate(lines):
        ls = line.strip()
        if ls == '## 要旨':
            section = 'ja_abstract'
            continue
        elif ls == '## Abstract':
            section = 'en_abstract'
            continue
        elif ls.startswith('## ') and section in ('ja_abstract', 'en_abstract', None):
            section = None
            continue
        elif ls == '---' and section in ('ja_abstract', 'en_abstract'):
            section = None
            continue
        if section == 'ja_abstract':
            if ls.startswith('**キーワード'):
                ja_keywords = re.sub(r'\*\*キーワード[：:]\*\*\s*', '', ls)
                section = None
            elif ls:
                ja_abstract_lines.append(ls)
        elif section == 'en_abstract':
            if ls.startswith('**Keywords'):
                en_keywords = re.sub(r'\*\*Keywords[：:]\*\*\s*', '', ls)
                section = None
            elif ls:
                en_abstract_lines.append(ls)

    # 和文要旨
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('要旨')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    if ja_abstract_lines:
        for aline in ja_abstract_lines:
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', aline).lstrip('　')
            p = doc.add_paragraph()
            run = p.add_run('　' + clean)
            run.font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        run = p.add_run('★要旨未作成')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 0, 0)

    if ja_keywords:
        p = doc.add_paragraph()
        run = p.add_run('キーワード：' + ja_keywords)
        run.font.size = Pt(10)

    # 英文要旨
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Abstract')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    if en_abstract_lines:
        for aline in en_abstract_lines:
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', aline)
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    else:
        p = doc.add_paragraph()
        run = p.add_run('★Abstract not yet written')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 0, 0)

    if en_keywords:
        p = doc.add_paragraph()
        run = p.add_run('Keywords: ' + en_keywords)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    doc.add_page_break()

    # 本文パース（緒言〜引用文献の手前まで + 引用文献）
    in_table = False
    table_rows = []
    table_header = []
    pending_figure = None
    inserted_figs = set()

    for i in range(body_start, len(lines)):
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            if in_table and table_rows:
                # テーブル終了 → docxテーブル生成
                _add_table(doc, table_header, table_rows)
                table_rows = []
                table_header = []
                in_table = False
            continue

        # 水平線
        if stripped == '---':
            continue

        # Markdownテーブル
        if stripped.startswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(set(c) <= set('-: ') for c in cells):
                # セパレータ行 → スキップ
                continue
            if not in_table:
                in_table = True
                table_header = cells
            else:
                table_rows.append(cells)
            continue

        # テーブルキャプション
        if stripped.startswith('**Table'):
            if in_table and table_rows:
                _add_table(doc, table_header, table_rows)
                table_rows = []
                table_header = []
                in_table = False
            clean = re.sub(r'\*\*', '', stripped)
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(10)
            continue

        # テーブル脚注（「妥当」は〜 で始まる行）
        if stripped.startswith('「') and in_table:
            if table_rows:
                _add_table(doc, table_header, table_rows)
                table_rows = []
                table_header = []
                in_table = False
            clean = _clean_md(stripped)
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.font.size = Pt(9)
            run.italic = True
            continue

        if in_table and table_rows:
            _add_table(doc, table_header, table_rows)
            table_rows = []
            table_header = []
            in_table = False

        # セクション見出し (## / ###)
        if stripped.startswith('## '):
            title_text = stripped.lstrip('# ').strip()
            p = doc.add_paragraph()
            p.space_before = Pt(18)
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Yu Gothic'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
            continue

        if stripped.startswith('### '):
            title_text = stripped.lstrip('# ').strip()
            p = doc.add_paragraph()
            p.space_before = Pt(12)
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Yu Gothic'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
            continue

        if stripped.startswith('#### '):
            title_text = stripped.lstrip('# ').strip()
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(11)
            continue

        # 通常段落
        clean = _clean_md(stripped)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(12)  # 一字下げ
        _add_formatted_text(p, stripped)

        # 図の挿入チェック（本文中にFigure Xへの言及がある段落の直後）
        if include_figures:
            for fig_key, (fig_file, fig_caption) in FIG_MAP.items():
                if fig_key in stripped or f'（{fig_key}）' in stripped:
                    fig_path = os.path.join(FIG_DIR, fig_file)
                    if os.path.exists(fig_path):
                        p_fig = doc.add_paragraph()
                        p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p_fig.add_run()
                        run.add_picture(fig_path, width=Mm(150))
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p_cap.add_run(fig_caption)
                        run.font.size = Pt(10)
                        run.italic = True
                        p_cap.space_after = Pt(12)
                        inserted_figs.add(fig_key)
                    break

        # Figure 4/5: セクション見出しで挿入（本文に明示的な (Figure X) がない場合）
        if include_figures:
            if '4. 時系列トレンド' in stripped and 'Figure 4' not in inserted_figs:
                pass  # Figure 1 already covers timeseries
            if '7. 論文タイプ別' in stripped:
                # Figure 5 を直後に挿入
                fig_path = os.path.join(FIG_DIR, FIG_MAP['Figure 5'][0])
                if os.path.exists(fig_path) and 'Figure 5' not in inserted_figs:
                    p_fig = doc.add_paragraph()
                    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_fig.add_run()
                    run.add_picture(fig_path, width=Mm(150))
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_cap.add_run(FIG_MAP['Figure 5'][1])
                    run.font.size = Pt(10)
                    run.italic = True
                    p_cap.space_after = Pt(12)
                    inserted_figs.add('Figure 5')

    # 残りのテーブル
    if in_table and table_rows:
        _add_table(doc, table_header, table_rows)

    # 引用文献の後の追加事項
    # COI
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run('利益相反（COI）')
    run.bold = True
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run('★小川先生確認後に記入')
    run.font.color.rgb = RGBColor(255, 0, 0)

    return doc


def _clean_md(text):
    """Markdown記法を除去"""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]', r'\1', text)
    return text


def _add_formatted_text(paragraph, md_text):
    """太字等のMarkdown書式をdocxのRunに変換"""
    # **太字** をパース
    parts = re.split(r'(\*\*[^*]+\*\*)', md_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Yu Mincho'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Mincho')
        else:
            # [X] 形式のリンクは除去
            clean = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', part)
            run = paragraph.add_run(clean)
            run.font.size = Pt(12)
            run.font.name = 'Yu Mincho'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Mincho')


def _add_table(doc, header, rows):
    """Markdownテーブル → docxテーブル"""
    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.autofit = True

    # ヘッダー行
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True
        run.font.size = Pt(9)

    # データ行
    for i, row in enumerate(rows):
        for j in range(min(len(row), ncols)):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', row[j])
            run = p.add_run(clean)
            run.font.size = Pt(9)

    doc.add_paragraph()  # 表後のスペース


if __name__ == '__main__':
    print('Building 01_manuscript.docx (with figures)...')
    doc = build_docx(include_figures=True)
    doc.save(OUT_FULL)
    print(f'  Saved: {OUT_FULL}')

    print('Building 02_manuscript_text_only.docx (no figures)...')
    doc2 = build_docx(include_figures=False)
    doc2.save(OUT_TEXT)
    print(f'  Saved: {OUT_TEXT}')

    # ファイルサイズ
    for f in [OUT_FULL, OUT_TEXT]:
        sz = os.path.getsize(f)
        print(f'  {os.path.basename(f)}: {sz//1024} KB')

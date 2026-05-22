#!/usr/bin/env python3
"""投稿用: 表docx個別ファイル + カバーレター + README"""

import os, re
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
SUB = os.path.join(BASE, 'submission')

def make_table_doc(title, header, rows, footnote=None, filename=None):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    style = doc.styles['Normal']
    style.font.size = Pt(10)
    style.font.name = 'Yu Mincho'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Mincho')

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)

    ncols = len(header)
    table = doc.add_table(rows=1+len(rows), cols=ncols)
    table.style = 'Table Grid'
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j in range(min(len(row), ncols)):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(re.sub(r'\*\*([^*]+)\*\*', r'\1', row[j]))
            run.font.size = Pt(9)

    if footnote:
        p = doc.add_paragraph()
        run = p.add_run(footnote)
        run.font.size = Pt(8)
        run.italic = True

    path = os.path.join(SUB, 'tables', filename)
    doc.save(path)
    print(f'  Saved: {path}')

# Table 1
make_table_doc(
    'Table 1. 認知ギャップの主要指標（Conservative辞書v2.1）',
    ['', 'kampo群', 'pubmed_kampo群'],
    [
        ['n（抄録保有）', '1,666', '4,446'],
        ['Q1（処方＋思考＋）', '529 (31.8%)', '34 (0.8%)'],
        ['Q2（処方−思考＋）', '162 (9.7%)', '328 (7.4%)'],
        ['Q3（処方＋思考−）', '577 (34.6%)', '1,341 (30.2%)'],
        ['Q4（処方−思考−）', '398 (23.9%)', '2,743 (61.7%)'],
        ['Q3/(Q1+Q3)', '52.2%', '97.5%'],
        ['OR (95% CI)', '', '36.16 (25.12-53.42)'],
        ['p値', '', '4.19 x 10^-177'],
    ],
    filename='Table1_quadrant.docx'
)

# Table 2
make_table_doc(
    'Table 2. 感度分析：3段階の辞書による認知ギャップ比率',
    ['source', 'Liberal', 'Conservative', 'Strict'],
    [
        ['kampo', '42.9%', '52.2%', '52.4%'],
        ['pubmed_kampo', '97.5%', '97.5%', '97.5%'],
    ],
    footnote='Liberal: 一般症候語を含む全語。Conservative: 偽陽性リスクの高い一般語（出血・発熱・浮腫・動悸・冷え）を除外。Strict: Conservativeからyin/yang deficiency等を追加除外。',
    filename='Table2_sensitivity.docx'
)

# Table 3
make_table_doc(
    'Table 3. サンプリング検証の結果（二次検証、計101件）',
    ['グループ', 'n', '妥当', '要注意', '保留', '誤分類'],
    [
        ['kampo Q1', '26', '19 (73%)', '5 (19%)', '2 (8%)', '0 (0%)'],
        ['kampo Q3', '25', '18 (72%)', '—', '6 (24%)', '1 (4%)'],
        ['pubmed Q1', '25', '8 (32%)', '12 (48%)', '3 (12%)', '2 (8%)'],
        ['pubmed Q3', '25', '25 (100%)', '0 (0%)', '0 (0%)', '0 (0%)'],
        ['合計', '101', '70 (69%)', '17 (17%)', '11 (11%)', '3 (3%)'],
    ],
    footnote='「妥当」は辞書の分類が人手評価と一致した論文。「要注意」は辞書がマッチしたが漢方思考概念が処方選択を駆動しているとは判断できない論文（定型句・背景説明・分類語彙の借用等）。「保留」は偽陰性候補または情報不足。「誤分類」は明確な誤分類。考察における偽陽性率の算出にあたっては「要注意」を偽陽性に含めた。',
    filename='Table3_validation.docx'
)

# カバーレター
print('Building cover letter...')
doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(30)
section.left_margin = Mm(30)
section.right_margin = Mm(30)
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Yu Mincho'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Mincho')
style.paragraph_format.line_spacing = Pt(22)

p = doc.add_paragraph('2026年3月31日')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph('')
doc.add_paragraph('日本東洋医学雑誌 編集委員会 御中')
doc.add_paragraph('')
doc.add_paragraph('拝啓')
doc.add_paragraph('')

body_text = (
    '下記の原著論文を日本東洋医学雑誌に投稿いたしたく、お願い申し上げます。'
)
doc.add_paragraph(body_text)
doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('タイトル：')
run.bold = True
p.add_run('漢方関連文献における処方名言及と漢方思考概念言及の乖離——10,535件のテキストマイニング')

doc.add_paragraph('')

desc = (
    '本論文は、漢方医学関連の学術文献10,535件を対象としたテキストマイニングにより、'
    '処方名への言及と漢方思考概念への言及の乖離（認知ギャップ）を初めて定量化した'
    '原著論文です。日本東洋医学雑誌では52.2%、PubMed漢方論文では97.5%の認知ギャップ'
    'が認められました。本研究の知見は、漢方医学の教育・研究・臨床の質の向上に向けた'
    '具体的な介入点を示すものと考えます。'
)
doc.add_paragraph(desc)
doc.add_paragraph('')

doc.add_paragraph('本論文は他誌に投稿中ではなく、著者全員が最終版を承認しています。')
doc.add_paragraph('すべての著者はICMJEの著者資格基準を満たしています。')
doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('利益相反：★小川先生確認後に記入')
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_paragraph('')
doc.add_paragraph('何卒ご査収のほど、よろしくお願い申し上げます。')
doc.add_paragraph('')
doc.add_paragraph('敬具')
doc.add_paragraph('')

p = doc.add_paragraph('責任著者：華岡晃生')
doc.add_paragraph('公立穴水総合病院内科 / 広島大学病院漢方診療センター')
p = doc.add_paragraph()
run = p.add_run('Email: ★')
run.font.color.rgb = RGBColor(255, 0, 0)
p = doc.add_paragraph()
run = p.add_run('Tel: ★')
run.font.color.rgb = RGBColor(255, 0, 0)

cover_path = os.path.join(SUB, 'cover_letter', 'cover_letter.docx')
doc.save(cover_path)
print(f'  Saved: {cover_path}')

print('Done.')
